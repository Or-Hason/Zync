"""Byte-level file type detection and raw text extraction.

Security note (MIME validation):
    File type is decided from the *bytes*, never the client-supplied extension.
    ``python-magic`` (libmagic) is the primary signal. PDFs are accepted on the
    ``application/pdf`` magic verdict. DOCX is an Office-OpenXML ZIP container
    that several libmagic builds (notably the Windows binaries) mislabel as
    ``application/zip`` or ``application/octet-stream``; for those we confirm the
    OOXML structure by inspecting the ZIP entries. A renamed ``.exe`` therefore
    fails both paths and is rejected.
"""

from __future__ import annotations

import io
import logging
import zipfile
from typing import Literal

import magic
from docx import Document
from pdfminer.high_level import extract_text as pdf_extract_text

logger = logging.getLogger(__name__)

FileKind = Literal["pdf", "docx"]

# Canonical MIME types accepted for upload.
MIME_PDF = "application/pdf"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# libmagic verdicts that may front a genuine DOCX container and therefore
# warrant a structural OOXML confirmation before acceptance.
_DOCX_CANDIDATE_MIMES = frozenset(
    {MIME_DOCX, "application/zip", "application/octet-stream"}
)

# A valid DOCX ZIP always contains the main document part.
_DOCX_REQUIRED_ENTRY = "word/document.xml"


class UnsupportedFileTypeError(Exception):
    """Raised when uploaded bytes are neither a valid PDF nor a valid DOCX."""


class TextExtractionError(Exception):
    """Raised when a supported file type yields no extractable text."""


def _looks_like_docx(data: bytes) -> bool:
    """Return ``True`` when bytes are an OOXML wordprocessing (DOCX) container.

    Args:
        data: Raw uploaded file bytes.

    Returns:
        ``True`` if the bytes form a ZIP archive containing
        ``word/document.xml``; ``False`` otherwise.
    """
    if not zipfile.is_zipfile(io.BytesIO(data)):
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            return _DOCX_REQUIRED_ENTRY in archive.namelist()
    except zipfile.BadZipFile:
        return False


def detect_file_kind(data: bytes) -> FileKind:
    """Determine the file kind from its bytes using libmagic plus OOXML checks.

    Args:
        data: Raw uploaded file bytes.

    Returns:
        ``"pdf"`` or ``"docx"``.

    Raises:
        UnsupportedFileTypeError: If the bytes are neither a PDF nor a DOCX.
    """
    detected_mime = magic.Magic(mime=True).from_buffer(data)

    if detected_mime == MIME_PDF:
        return "pdf"

    if detected_mime in _DOCX_CANDIDATE_MIMES and _looks_like_docx(data):
        return "docx"

    logger.warning(
        "Rejected upload with disallowed MIME", extra={"mime": detected_mime}
    )
    raise UnsupportedFileTypeError(detected_mime)


def _extract_pdf_text(data: bytes) -> str:
    """Extract text from PDF bytes via pdfminer.six.

    Args:
        data: Raw PDF bytes.

    Returns:
        Extracted text (may be empty for image-only PDFs).
    """
    return pdf_extract_text(io.BytesIO(data)) or ""


def _extract_docx_text(data: bytes) -> str:
    """Extract text from DOCX bytes, including table cell contents.

    Args:
        data: Raw DOCX bytes.

    Returns:
        Newline-joined document text.
    """
    document = Document(io.BytesIO(data))
    lines: list[str] = [para.text for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_text(data: bytes, kind: FileKind) -> str:
    """Extract raw text from supported file bytes.

    Args:
        data: Raw uploaded file bytes.
        kind: The detected file kind from :func:`detect_file_kind`.

    Returns:
        Non-empty extracted text.

    Raises:
        TextExtractionError: If extraction fails or yields no usable text.
    """
    try:
        text = _extract_pdf_text(data) if kind == "pdf" else _extract_docx_text(data)
    except Exception as exc:  # noqa: BLE001 - normalise any parser failure.
        logger.error("Text extraction failed", extra={"kind": kind, "error": str(exc)})
        raise TextExtractionError(kind) from exc

    if not text.strip():
        raise TextExtractionError(kind)

    return text
